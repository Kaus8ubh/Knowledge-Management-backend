from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup  
import io
from weasyprint import HTML
from docx import Document
from html2docx import html2docx

class DocumentGenerator:

    def _convert_html_to_text(self, html_content):
        """Convert HTML content into formatted text for FPDF and python-docx."""
        # Use BeautifulSoup to parse and clean HTML content
        soup = BeautifulSoup(html_content, "html.parser")
        return soup

    def generate_card_pdf(self, card_data):
        """Generate a PDF with proper formatting using HTML and WeasyPrint."""

        # Format creation date
        created_at_str = ""
        if "created_at" in card_data:
            created_at = card_data["created_at"]
            if isinstance(created_at, str):
                try:
                    created_at = datetime.fromisoformat(created_at)
                except ValueError:
                    try:
                        created_at = datetime.strptime(created_at, "%Y-%m-%d %H:%M:%S")
                    except Exception:
                        created_at = None
            if created_at:
                created_at_str = f"<h3><em>Created: {created_at.strftime('%Y-%m-%d')}</em></h3>"

        # Build HTML content
        html_parts = [f"<h1>{card_data.get('title', 'Knowledge Card')}</h1>"]
        html_parts.append(created_at_str)

        if "source_url" in card_data:
            html_parts.append(f"<h3 class='source'><em>Source: {card_data.get('source_url', 'N/A')}</em></h3>")

        # Optional fields section
        optional_fields = [
            ("Note", "note"),
            ("Summary", "summary"),
            ("QnA", "qna"),
            ("Knowledge Map", "knowledge_map")
        ]
        
        for title, key in optional_fields:
            if key in card_data:
                html_parts.append(f"<h2>{title}:</h2>")
                if key == "qna":
                    # Formatting QnA section
                    for qna in card_data[key]:
                        html_parts.append(f"<h3>Q: {qna['question']}</h3>")
                        html_parts.append(f"<p><strong>A:</strong> {qna['answer']}</p>")
                elif key == "knowledge_map":
                    # Formatting Knowledge Map section
                    for section in card_data[key]:
                        html_parts.append(f"<h3>{section['icon']} {section['section']}:</h3>")
                        for item in section['items']:
                            html_parts.append(f"<p><strong>Topic:</strong> {item['topic']}</p>")
                            html_parts.append(f"<p><strong>Description:</strong> {item['description']}</p>")
                            html_parts.append(f"<p><strong>Difficulty:</strong> {item['difficulty']}</p>")
                            html_parts.append("<hr>")  # Add a horizontal line to separate items
                else:
                    # Simple text for other fields
                    html_parts.append(f"<p>{card_data[key]}</p>")

        # Join the HTML content together
        html_content = "<html><head><style>body {font-family: Arial, sans-serif; padding: 20px;} h1 {font-size: 28px; font-weight: bold;} h2 {font-size: 18px;} h3 {font-size: 16px; margin-top: 20px;} p {font-size: 14px; line-height: 1.6;} .source {color: blue;} hr {border: 1px solid #ccc; margin-top: 10px;}</style></head><body>" + "".join(html_parts) + "</body></html>"

        # Generate and return PDF as bytes using WeasyPrint
        pdf_bytes = HTML(string=html_content).write_pdf()
        return pdf_bytes
    
    def generate_card_docx(self, card_data):
        """Generate a DOCX file with properly formatted QnA and Knowledge Map sections."""
        try:
            # Create a new Document
            doc = Document()

            # Add title and creation date
            doc.add_heading(card_data.get('title', 'Knowledge Card'), level=1)

            created_at_str = ""
            if "created_at" in card_data:
                created_at = card_data["created_at"]
                if isinstance(created_at, str):
                    try:
                        created_at = datetime.fromisoformat(created_at)
                    except ValueError:
                        try:
                            created_at = datetime.strptime(created_at, "%Y-%m-%d %H:%M:%S")
                        except Exception:
                            created_at = None
                if created_at:
                    created_at_str = f"Created: {created_at.strftime('%Y-%m-%d')}"
                    doc.add_paragraph(created_at_str)

            if "source_url" in card_data:
                doc.add_paragraph(f"Source: {card_data.get('source_url', 'N/A')}")

            # Optional fields section
            optional_fields = [
                ("Note", "note"),
                ("Summary", "summary"),
                ("QnA", "qna"),
                ("Knowledge Map", "knowledge_map")
            ]

            for title, key in optional_fields:
                if key in card_data:
                    # Add title for each section (e.g., QnA, Knowledge Map)
                    doc.add_heading(title, level=2)

                    if key == "qna":
                        # Formatting QnA section
                        for qna in card_data[key]:
                            doc.add_paragraph(f"Q: {qna['question']}", style='List Number')
                            doc.add_paragraph(f"A: {qna['answer']}", style='BodyText')
                            doc.add_paragraph()  # Add a blank line between Q&A pairs

                    elif key == "knowledge_map":
                        # Formatting Knowledge Map section
                        for section in card_data[key]:
                            doc.add_heading(f"{section['icon']} {section['section']}", level=3)

                            for item in section['items']:
                                doc.add_paragraph(f"Topic: {item['topic']}", style='List Bullet')
                                doc.add_paragraph(f"Description: {item['description']}", style='BodyText')
                                doc.add_paragraph(f"Difficulty: {item['difficulty']}", style='BodyText')
                                doc.add_paragraph()  # Add a blank line between items

                    else:
                        # Add simple text for other fields
                        doc.add_paragraph(card_data[key], style='BodyText')

            # Save the DOCX file with a unique name
            file_name = "knowledge_card_v2.docx"
            doc.save(file_name)

            print(f"Document saved as {file_name} successfully!")

        except Exception as e:
            print(f"An error occurred while generating the DOCX: {e}")